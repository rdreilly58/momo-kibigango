#!/usr/bin/env python3
"""
docx.py — Read and write Microsoft Word (.docx) files

Usage:
  docx.py read FILE.docx                    # Extract all text
  docx.py read FILE.docx --json             # JSON output with structure
  docx.py create "Title" -o FILE.docx       # Create new document
  docx.py append FILE.docx "text"           # Append paragraph
  docx.py append FILE.docx --heading "text" # Append heading
  docx.py append FILE.docx --list "item"    # Add list item
"""

import sys
import json
from pathlib import Path
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


def read_docx(filepath, as_json=False):
    """Read and extract text from a Word document."""
    doc = Document(filepath)
    
    if as_json:
        # Structured output
        data = {
            "document": filepath,
            "paragraphs": [],
            "tables": []
        }
        
        for para in doc.paragraphs:
            if para.text.strip():
                data["paragraphs"].append({
                    "text": para.text,
                    "level": para.style.base_style.name if para.style else None,
                    "style": para.style.name if para.style else "Normal"
                })
        
        for i, table in enumerate(doc.tables):
            table_data = {
                "table_number": i,
                "rows": []
            }
            for row in table.rows:
                table_data["rows"].append([cell.text for cell in row.cells])
            data["tables"].append(table_data)
        
        return json.dumps(data, indent=2)
    else:
        # Plain text output
        lines = []
        for para in doc.paragraphs:
            lines.append(para.text)
        for table in doc.tables:
            lines.append("\n[TABLE]")
            for row in table.rows:
                lines.append(" | ".join(cell.text for cell in row.cells))
            lines.append("[END TABLE]\n")
        
        return "\n".join(lines)


def create_docx(title, filepath, author=None):
    """Create a new Word document."""
    doc = Document()
    
    if title:
        heading = doc.add_heading(title, level=1)
    
    if author:
        doc.core_properties.author = author
    
    doc.save(filepath)
    print(f"[docx] Created: {filepath}")


def append_to_docx(filepath, text, style="Normal", heading_level=None, is_list=False):
    """Append content to an existing Word document."""
    doc = Document(filepath)
    
    if heading_level:
        doc.add_heading(text, level=heading_level)
    elif is_list:
        doc.add_paragraph(text, style="List Bullet")
    else:
        doc.add_paragraph(text, style=style)
    
    doc.save(filepath)
    print(f"[docx] Updated: {filepath}")


def main():
    if len(sys.argv) < 2:
        print(__doc__)
        sys.exit(1)
    
    cmd = sys.argv[1]
    
    if cmd == "read":
        filepath = sys.argv[2]
        as_json = "--json" in sys.argv
        result = read_docx(filepath, as_json=as_json)
        print(result)
    
    elif cmd == "create":
        title = sys.argv[2] if len(sys.argv) > 2 else ""
        output = None
        for i, arg in enumerate(sys.argv):
            if arg == "-o" and i + 1 < len(sys.argv):
                output = sys.argv[i + 1]
        
        if not output:
            print("Error: -o OUTPUT required")
            sys.exit(1)
        
        create_docx(title, output)
    
    elif cmd == "append":
        filepath = sys.argv[2]
        
        if not Path(filepath).exists():
            print(f"Error: file not found: {filepath}")
            sys.exit(1)
        
        text = None
        heading_level = None
        is_list = False
        
        for i, arg in enumerate(sys.argv[3:], 3):
            if arg == "--heading" and i + 1 < len(sys.argv):
                heading_level = int(sys.argv[i + 1])
                text = sys.argv[i + 1]
                break
            elif arg == "--list":
                is_list = True
                text = sys.argv[i + 1] if i + 1 < len(sys.argv) else ""
                break
            elif not arg.startswith("--"):
                text = arg
                break
        
        if not text:
            print("Error: no text provided")
            sys.exit(1)
        
        append_to_docx(filepath, text, heading_level=heading_level, is_list=is_list)
    
    else:
        print(f"Unknown command: {cmd}")
        print(__doc__)
        sys.exit(1)


if __name__ == "__main__":
    main()
